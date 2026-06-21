import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins (single column, ATS-safe) ──────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(1.6)
section.bottom_margin = Cm(1.6)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)

# ── Colour palette ───────────────────────────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
ACCENT = RGBColor(0x0F, 0x3D, 0x91)   # professional blue
GRAY   = RGBColor(0x55, 0x55, 0x55)

# ── Helper: remove space before/after paragraph ──────────────────────────────
def tight(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    return para

def add_run(para, text, bold=False, italic=False, size=10, color=DARK):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return run

# ── Horizontal rule helper ───────────────────────────────────────────────────
def add_hr(doc, color="0F3D91"):
    para = doc.add_paragraph()
    tight(para)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

# ── Section heading ──────────────────────────────────────────────────────────
def section_heading(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)
    add_run(para, title.upper(), bold=True, size=10.5, color=ACCENT)
    add_hr(doc)

# ── Job block ────────────────────────────────────────────────────────────────
def job_header(doc, company, title, period, location):
    # Company | Title
    p1 = doc.add_paragraph()
    tight(p1)
    p1.paragraph_format.space_before = Pt(5)
    add_run(p1, company, bold=True, size=10.5, color=DARK)
    add_run(p1, "  |  ", size=10, color=GRAY)
    add_run(p1, title, bold=True, size=10, color=ACCENT)
    # Period & location
    p2 = doc.add_paragraph()
    tight(p2)
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, f"{period}  ·  {location}", italic=True, size=9, color=GRAY)

def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    tight(para)
    para.paragraph_format.left_indent        = Inches(0.15)
    para.paragraph_format.first_line_indent  = Inches(-0.15)
    para.paragraph_format.space_after        = Pt(1.5)
    # split on ** for bold segments
    parts = text.split("**")
    for i, part in enumerate(parts):
        add_run(para, part, bold=(i % 2 == 1), size=9.5, color=DARK)
    return para

# ════════════════════════════════════════════════════════════════════════════
#  NAME & CONTACT
# ════════════════════════════════════════════════════════════════════════════
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(name_para)
name_para.paragraph_format.space_after = Pt(2)
add_run(name_para, "MUJTABA SAJAWAL", bold=True, size=20, color=DARK)

contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(contact_para)
contact_para.paragraph_format.space_after = Pt(2)
add_run(contact_para,
        "+971 5 432 27639  ·  mujtabasajawal@hotmail.com  ·  Dubai, UAE  ·  UAE Driving License",
        size=9.5, color=GRAY)

linkedin_para = doc.add_paragraph()
linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(linkedin_para)
linkedin_para.paragraph_format.space_after = Pt(3)
add_run(linkedin_para, "linkedin.com/in/mujtaba-sajawal", size=9.5, color=ACCENT)

headline_para = doc.add_paragraph()
headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(headline_para)
headline_para.paragraph_format.space_after = Pt(2)
add_run(headline_para,
        "Performance Marketing & AI Analytics Lead  |  10+ Years Experience",
        bold=True, size=10, color=DARK)

add_hr(doc, "1A1A2E")

# ════════════════════════════════════════════════════════════════════════════
#  SUMMARY
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Professional Summary")
summary = doc.add_paragraph()
tight(summary)
summary.paragraph_format.space_after = Pt(4)
add_run(summary,
    "Results-driven Performance Marketing & AI Analytics Lead with 10+ years of experience sitting at the intersection of marketing measurement, business intelligence, and AI enablement. "
    "Specializes in owning end-to-end attribution frameworks, establishing single-source-of-truth (SSOT) reporting systems, and driving organization-wide AI transformation. "
    "Proven builder and operator with hands-on experience engineering autonomous AI agents, custom scraper pipelines, locally-hosted CRM dashboards, and automated report pipelines (Python, APIs, webhooks, Power Automate, Make) that reduce operational effort by 40%. "
    "Managed budgets up to AED 35M and bridged advanced analytics with practical AI implementations to optimize CAC and maximize marketing efficiency.",
    size=9.5, color=DARK)

# ════════════════════════════════════════════════════════════════════════════
#  EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Professional Experience")

# ── 1. Cushman & Wakefield ───────────────────────────────────────────────────
job_header(doc, "Cushman & Wakefield | Core", "Performance Marketing Lead",
           "August 2025 – Present", "Dubai, UAE")
cw_bullets = [
    "**Paid Campaigns & ROI:** Managed B2B & B2C paid media across Google, Meta, and LinkedIn, driving a **34% increase in qualified lead volume** while reducing CPL by **18%**.",
    "**Attribution & Pipelines:** Implemented automated data pipelines and tracking scripts to feed performance metrics into a centralized marketing database, ensuring **100% data consistency**.",
    "**Power BI & Automation:** Developed Power BI tracking dashboards and automated weekly reporting pipelines using custom Python logic, saving **12+ hours/week** in reporting overhead.",
    "**AI-Enabled Leaderboard:** Replaced legacy Spinify leaderboard with a custom, **AI-built webpage** hosted locally and synced with **Salesforce CRM** to update sales metrics in real-time.",
    "**AI Development (IDEs):** Utilised **VS Code (Claude)** and **Google Antigravity (Gemini)** using both IDEs to build high-converting property landing pages, driving a **25% increase in inquiries**.",
    "**Programmatic Media (MCPs):** Activated Model Context Protocol (**MCPs**) for Google and Meta Ads, enabling AI-driven programmatic tracking and automated campaign adjustments.",
    "**CRM Database Migration:** Led Salesforce CRM operations and database migration to Microsoft Dynamics 365 Sales with **zero data loss** and seamless team adoption.",
]
for b in cw_bullets:
    bullet(doc, b)

# ── 2. RNS Realty ────────────────────────────────────────────────────────────
job_header(doc, "RNS Realty", "Digital Marketing Manager",
           "February 2025 – August 2025 (Contract)", "Dubai, UAE")
rns_bullets = [
    "**Lead Generation:** Spearheaded marketing initiatives across paid and organic channels, driving a **32% increase in overall lead volume** within 3 months.",
    "**ROI Optimization:** Managed integrated campaign budgets, optimizing spend allocation across search and social to achieve a **26% improvement in ROI**.",
    "**CRM & Retention:** Executed segmented email, WhatsApp, and CRM campaigns, achieving a **35% open rate** and a **17% conversion rate** on retargeted audiences.",
]
for b in rns_bullets:
    bullet(doc, b)

# ── 3. Knight Frank MENA ─────────────────────────────────────────────────────
job_header(doc, "Knight Frank MENA", "Performance Marketing Lead",
           "September 2023 – April 2025", "Dubai, UAE")
kf_bullets = [
    "**Campaign Growth:** Managed paid media budgets across search, social, and display (Google, Meta, LinkedIn, TikTok), boosting lead generation by **31%** and ROI by **2x**.",
    "**Attribution & ROI:** Drove a **12% uplift in overall ROI** by establishing a multi-touch attribution framework and reallocating budgets to high-performing channels.",
    "**A/B Testing:** Optimized acquisition funnels through continuous A/B testing, resulting in a **29% improvement in campaign ROI**.",
    "**Lead Nurturing:** Maintained an **85% lead-qualification rate** and reactivated **7% of dormant leads** through targeted CRM, email, and WhatsApp workflows.",
]
for b in kf_bullets:
    bullet(doc, b)

# ── 4. Footprint Real Estate ──────────────────────────────────────────────────
job_header(doc, "Footprint Real Estate", "Digital Marketing Specialist",
           "February 2023 – September 2023", "Dubai, UAE")
fp_bullets = [
    "**Multi-Channel Paid Acquisition:** Managed high-budget digital campaigns across Google Ads, Meta, LinkedIn, and TikTok, increasing website traffic by **40%** and app downloads by **53%**.",
    "**Performance Measurement:** Analyzed multi-channel campaign analytics to optimize conversion paths and align marketing spend with sales team outcomes.",
]
for b in fp_bullets:
    bullet(doc, b)

# ── 5. OLX Group ──────────────────────────────────────────────────────────────
job_header(doc, "OLX Group", "Category Manager",
           "April 2018 – January 2023", "Lahore, Pakistan")
olx_bullets = [
    "**Revenue Operations:** Delivered USD 120K yearly revenue target with **5% YoY growth**, managing category marketing campaigns.",
    "**Team Leadership:** Led a cross-functional team of 50+ members across 6 cities, earning the 'Employee of the Year' award in 2020.",
    "**Lifecycle Marketing:** Analyzed consumer offerings and acquisition channels, driving a **5% increase in traffic** and **8% boost in customer satisfaction**.",
]
for b in olx_bullets:
    bullet(doc, b)

# ════════════════════════════════════════════════════════════════════════════
#  SKILLS
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Core Skills")

skills = [
    ("Marketing Analytics & BI", "Single Source of Truth (SSOT) · Looker Studio · Power BI · Tableau · Funnel Analysis · Cohort Analysis · CAC & LTV Optimization · Incrementality Testing · Advanced Excel & Sheets"),
    ("Attribution & Tracking",  "UTM Governance · Event Tagging & Taxonomies · Google Tag Manager · Conversion API (CAPI) · Mobile Measurement Partners (MMP)"),
    ("AI & Automation",         "AI Agents & Workflows · VS Code & Google Antigravity (IDEs) · Custom AI Scrapers · Marketing Automation · Webhooks · API Integrations · Python · Power Automate · Make · Zapier · n8n"),
    ("CRMs & Databases",        "Salesforce · Microsoft Dynamics 365 · HubSpot · Zoho · SQL (Basic)"),
    ("Paid Media Platforms",    "Google Ads · Meta Ads · LinkedIn Ads · TikTok Ads · Display & Video 360 · Programmatic Buying"),
    ("Productivity & Collaboration", "Google Workspace · Microsoft Excel (VBA / PowerQuery) · Slack · Jira · Confluence"),
    ("Soft Skills",             "Budget Management (up to AED 35M) · Team Leadership · Cross-functional Collaboration · Problem Solving"),
]

for label, items in skills:
    p = doc.add_paragraph()
    tight(p)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{label}: ", bold=True, size=9.5, color=DARK)
    add_run(p, items, size=9.5, color=DARK)

# ════════════════════════════════════════════════════════════════════════════
#  CERTIFICATIONS
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Certifications")

certs = [
    ("TikTok Media Buying",              "TikTok",   "05/2023 – Present"),
    ("LinkedIn Ads",                     "LinkedIn", "01/2023 – Present"),
    ("X Ads Manager",                    "X",        "08/2022 – Present"),
    ("Google Ads Apps Certification",    "Google",   "07/2022 – Present"),
    ("PPC Automation",                   "SEMrush",  "07/2022 – Present"),
    ("Display & Video 360",              "Google",   "05/2021 – Present"),
    ("Meta Certification",               "META",     "09/2020 – Present"),
    ("Google Analytics Certification",   "Google",   "03/2020 – Present"),
]

for name, issuer, period in certs:
    p = doc.add_paragraph()
    tight(p)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{name}", bold=True, size=9.5, color=DARK)
    detail = f"  ·  {issuer}"
    if period:
        detail += f"  ·  {period}"
    add_run(p, detail, size=9.5, color=GRAY)

# ════════════════════════════════════════════════════════════════════════════
#  EDUCATION
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Education")

edu = [
    ("Bachelor of Business & Economics", "University of Punjab", "2007 – 2010"),
]

for degree, school, period in edu:
    p = doc.add_paragraph()
    tight(p)
    p.paragraph_format.space_after = Pt(2.5)
    add_run(p, degree, bold=True, size=9.5, color=DARK)
    add_run(p, f"  ·  {school}  ·  {period}", size=9.5, color=GRAY)

# ════════════════════════════════════════════════════════════════════════════
#  LANGUAGES
# ════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Languages")
lang_p = doc.add_paragraph()
tight(lang_p)
add_run(lang_p, "English", bold=True, size=9.5, color=DARK)
add_run(lang_p, " (Fluent)  ·  ", size=9.5, color=GRAY)
add_run(lang_p, "Urdu", bold=True, size=9.5, color=DARK)
add_run(lang_p, " (Native)  ·  ", size=9.5, color=GRAY)
add_run(lang_p, "Punjabi", bold=True, size=9.5, color=DARK)
add_run(lang_p, " (Native)  ·  ", size=9.5, color=GRAY)
add_run(lang_p, "Hindi", bold=True, size=9.5, color=DARK)
add_run(lang_p, " (Intermediate)", size=9.5, color=GRAY)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "/Users/mujtabasajawal/Downloads/Resume/Mujtaba_Sajawal_Marketing_AI_Analytics_Lead_Resume.docx"
doc.save(output_path)
print(f"Tailored resume saved successfully to: {output_path}")
